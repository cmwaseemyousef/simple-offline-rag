import re
import sys
from pathlib import Path

from markdown import markdown
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
BLOG_MD = ROOT / "blog.md"
BLOG_HTML = ROOT / "blog.html"
BLOG_DOCX = ROOT / "blog.docx"


def strip_frontmatter(text: str):
    # Returns (meta dict, body)
    if text.startswith("---\n"):
        end = text.find("\n---\n", 4)
        if end != -1:
            fm = text[4:end].strip()
            body = text[end + 5 :]
            meta = {}
            for line in fm.splitlines():
                if ":" in line:
                    k, v = line.split(":", 1)
                    meta[k.strip()] = v.strip()
            return meta, body
    return {}, text


def to_html(md_text: str) -> str:
    return markdown(md_text, extensions=["fenced_code", "tables", "toc", "codehilite"])  # type: ignore


def add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    run = p.add_run(code)
    font = run.font
    font.name = "Consolas"
    # Ensure font mapping for Word
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    font.size = Pt(10)


def md_to_docx(doc: Document, md_body: str):
    lines = md_body.splitlines()
    in_code = False
    code_buf = []
    for raw in lines:
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_buf))
            continue
        if in_code:
            code_buf.append(line)
            continue
        # Headings
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
        elif line.startswith("- "):
            # simple bullet list
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("*") and line.lstrip().startswith("* "):
            p = doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)


def main():
    if not BLOG_MD.exists():
        print(f"blog.md not found at {BLOG_MD}")
        sys.exit(1)
    raw = BLOG_MD.read_text(encoding="utf-8")
    meta, body = strip_frontmatter(raw)

    # HTML
    html = to_html(body)
    BLOG_HTML.write_text(html, encoding="utf-8")
    print(f"Wrote {BLOG_HTML}")

    # DOCX
    doc = Document()
    title = meta.get("title") or "Blog"
    desc = meta.get("description") or ""
    author = meta.get("author") or ""

    if title:
        add_heading(doc, title, level=0)
    if desc:
        p = doc.add_paragraph(desc)
        p.runs[0].italic = True
    if author:
        doc.add_paragraph(f"Author: {author}")

    md_to_docx(doc, body)
    doc.save(str(BLOG_DOCX))
    print(f"Wrote {BLOG_DOCX}")


if __name__ == "__main__":
    main()
